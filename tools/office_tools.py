"""
Jarvis Tools — Microsoft Office Automation
===========================================
Defines the tool interfaces for reading and writing Excel spreadsheets (.xlsx)
and editing Word documents (.docx) using native Python libraries.
"""

import os
from pathlib import Path
from loguru import logger
import openpyxl
from docx import Document


async def modify_excel_cell(file_path: str, sheet_name: str, cell_name: str, value: str) -> str:
    """
    Write or update a cell value in an Excel spreadsheet (.xlsx).
    If the Excel file does not exist, a new spreadsheet will be created.

    Args:
        file_path: The absolute path of the Excel file (e.g., "C:/Users/hp/Documents/sales.xlsx").
        sheet_name: The sheet tab name (e.g., "Sheet1", "Monthly Data").
        cell_name: The cell identifier (e.g., "A1", "C5").
        value: The text or numeric content to write into the cell.
    """
    logger.info(f"🔧 Tool Call: Writing value '{value}' to Excel cell '{cell_name}' in sheet '{sheet_name}'...")
    
    resolved_path = Path(os.path.expandvars(file_path)).resolve()
    
    try:
        # Load or create workbook
        if resolved_path.exists():
            wb = openpyxl.load_workbook(resolved_path)
        else:
            # Ensure parent directories exist
            resolved_path.parent.mkdir(parents=True, exist_ok=True)
            wb = openpyxl.Workbook()
            
        # Get or create sheet
        if sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
        else:
            # If default sheet is empty and named "Sheet", rename it
            if len(wb.sheetnames) == 1 and wb.sheetnames[0] == "Sheet":
                sheet = wb.active
                sheet.title = sheet_name
            else:
                sheet = wb.create_sheet(title=sheet_name)
                
        # Set cell value (convert to float/int if possible)
        try:
            if value.isdigit():
                sheet[cell_name] = int(value)
            elif value.replace('.', '', 1).isdigit() and value.count('.') == 1:
                sheet[cell_name] = float(value)
            else:
                sheet[cell_name] = value
        except Exception:
            sheet[cell_name] = value
            
        wb.save(resolved_path)
        wb.close()
        
        return f"Successfully updated Excel spreadsheet at '{resolved_path}' on Sheet '{sheet_name}' at Cell '{cell_name}'."
    except Exception as e:
        logger.error(f"Excel error: {e}")
        return f"Failed to modify Excel cell. Error: {e}"


async def read_excel_cell(file_path: str, sheet_name: str, cell_name: str) -> str:
    """
    Read the value of a cell from an Excel spreadsheet (.xlsx).

    Args:
        file_path: The absolute path of the Excel file.
        sheet_name: The sheet tab name (e.g. "Sheet1").
        cell_name: The cell coordinate (e.g., "B2").
    """
    logger.info(f"🔧 Tool Call: Reading Excel cell '{cell_name}' in sheet '{sheet_name}'...")
    
    resolved_path = Path(os.path.expandvars(file_path)).resolve()
    
    if not resolved_path.exists():
        return f"File '{resolved_path}' does not exist."
        
    try:
        wb = openpyxl.load_workbook(resolved_path, read_only=True)
        if sheet_name not in wb.sheetnames:
            return f"Sheet '{sheet_name}' not found in workbook."
            
        sheet = wb[sheet_name]
        val = sheet[cell_name].value
        wb.close()
        
        return f"The value of cell {cell_name} in Sheet '{sheet_name}' is: '{val}'."
    except Exception as e:
        logger.error(f"Excel read error: {e}")
        return f"Failed to read Excel cell. Error: {e}"


async def modify_word_document(file_path: str, search_text: str, replace_text: str) -> str:
    """
    Search for a text string and replace it with a new string in a Word document (.docx).

    Args:
        file_path: The absolute path of the Word document.
        search_text: The exact text string to search for.
        replace_text: The new text string to replace the search text.
    """
    logger.info(f"🔧 Tool Call: Replacing '{search_text}' with '{replace_text}' in Word document '{file_path}'...")
    
    resolved_path = Path(os.path.expandvars(file_path)).resolve()
    if not resolved_path.exists():
        return f"The Word document at '{resolved_path}' does not exist."
        
    try:
        doc = Document(resolved_path)
        replace_count = 0
        
        # 1. Search and replace in paragraphs
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                paragraph.text = paragraph.text.replace(search_text, replace_text)
                replace_count += 1
                
        # 2. Search and replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if search_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(search_text, replace_text)
                            replace_count += 1
                            
        doc.save(resolved_path)
        return f"Successfully updated Word document. Replaced {replace_count} occurrences of '{search_text}' with '{replace_text}'."
    except Exception as e:
        logger.error(f"Word document edit error: {e}")
        return f"Failed to edit Word document. Error: {e}"
